from django.shortcuts import render
from django.contrib.auth import authenticate, logout, login
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, HttpResponsePermanentRedirect, JsonResponse
from django.urls import reverse
from django.utils import timezone
from django.views.decorators.csrf import csrf_exempt

from .models import User, Poll_choice, Poll, User_poll_choice, Event, Event_register, Article, Article_Image, Document

import datetime
import os
import shutil
import requests
import pyqrcode
import json
import pytz
from docx import Document as docxdoc
from docx.shared import Inches
import openpyxl
from openpyxl.styles import Font

# TZ = pytz.timezone('Europe/Moscow')
STATIC_PATH = os.getcwd()+"/joseph_app/static"
FILE_PATH = "/joseph_app/"

def joseph_index(request):
    return HttpResponsePermanentRedirect(reverse("joseph_app:index"))

def index(request):
    article_list = []
    for article in Article.objects.order_by("-date")[:3]:
        try:
            image = article.article_image_set.earliest("pk")
        except:
            image = None
        to_send = {
            "article" : article,
            "image" : image,
        }
        article_list.append(to_send)
    event_list = []
    for event in Event.objects.filter(event_date__gte=timezone.now()).order_by("event_date")[:3]:
        to_send = {
            "event" : event,
        }
        event_list.append(to_send)
    response = {
        "news" : article_list,
        "events" : event_list,
    }
    return render(request, 'joseph_app/index.html', response)

#Логин пользователя, редирект на user_cab
def login_user(request):
    email = request.POST["email"]
    password = request.POST["password"]
    user = authenticate(request, username=email, password=password)
    if user is not None:
        login(request, user)
        return HttpResponsePermanentRedirect(reverse("joseph_app:user_cab", args=(user.pk,)))
    else:
        return HttpResponse("Login failed")

def logout_user(request):
    logout(request)
    return HttpResponsePermanentRedirect(reverse("joseph_app:index"))

#Создание нового пользователя, рендер страницы регистрации update_user.html
def create_user(request):
    return render(request, "joseph_app/register.html")

#Регистрация нового пользователя в базе данных, редирект на login_user
def register_user(request):
    user = User.objects.create_user(email=request.POST['email_new'], password=request.POST['password_new'])
    user.name = request.POST['name']
    user.surname = request.POST['surname']
    user.second_name = request.POST['second_name']
    if request.FILES:
        f = request.FILES['avatar']
        path = STATIC_PATH+FILE_PATH+"userdata/{}/avatar/".format(user.pk)
        if not os.path.exists(path):
            os.mkdir(path[:-8])
            os.mkdir(path[:-1])
        with open(path+f.name, "wb+") as destination:
            for chunk in f.chunks():
                destination.write(chunk)
        user.avatar = path[len(STATIC_PATH):]+f.name
    if request.POST['date_of_birth']:
        dob = request.POST['date_of_birth']
        user.date_of_birth = dob[6:]+"-"+dob[3:5]+"-"+dob[:2]
    user.phone = request.POST['phone']
    user.pubnet = request.POST['pubnet']
    user.university = request.POST['university']
    user.course = request.POST['course']
    user.reg_address = request.POST['reg_address']
    user.cur_address = request.POST['cur_address']
    user.organizations = request.POST['organizations']
    user.hobby = request.POST['hobby']
    user.save()
    return HttpResponsePermanentRedirect(reverse("joseph_app:login"))

#Изменение информации о пользователе, рендер update_user.html ???
def redact_user(request, user_pk):
    user = request.user
    if user is not None and user.is_authenticated and user.pk == user_pk:
        response = {
            "user" : user,
        }
        return render(request, "registration/update_user.html", response)

#Изменение информации о пользователе в базе данных, редирект на user_cab
def update_user(request, user_pk):
    user = request.user
    if user is not None and user.is_authenticated and user.pk == user_pk:
        user.name = request.POST['name']
        user.surname = request.POST['surname']
        user.second_name = request.POST['second_name']
        if request.FILES:
            f = request.FILES['avatar']
            path = STATIC_PATH+FILE_PATH+"userdata/{}/avatar/".format(user.pk)
            if not os.path.exists(path):
                os.mkdir(path[:-8])
                os.mkdir(path[:-1])
            else:
                os.remove(user.avatar)
            with open(path + f.name, "wb+") as destination:
                for chunk in f.chunks():
                    destination.write(chunk)
            user.avatar = path[len(STATIC_PATH):]+f.name
        dob = request.POST['date_of_birth']
        user.date_of_birth = dob[6:]+"-"+dob[3:5]+"-"+dob[:2]
        user.phone = request.POST['phone']
        user.pubnet = request.POST['pubnet']
        user.university = request.POST['university']
        user.course = request.POST['course']
        user.reg_address = request.POST['reg_address']
        user.cur_address = request.POST['cur_address']
        user.organizations = request.POST['organizations']
        user.hobby = request.POST['hobby']
        user.save()
        return HttpResponsePermanentRedirect(reverse("joseph_app:user_cab", args=(user.pk,)))

def password_change(request, user_pk):
    pwd = request.POST['password_old']
    email = request.user.email
    user = authenticate(request, username=email, password=pwd)
    if user is not None and user.is_authenticated:
        pwd1 = request.POST['password_new']
        pwd2 = request.POST['password_val']
        if pwd1 == pwd2:
            logout(request)
            user.set_password(pwd1)
            user.save()
        return HttpResponsePermanentRedirect(reverse("joseph_app:login"))

#Рендер страницы личного кабинета cabinet.html
@login_required(login_url="/joseph")
def user_cab(request, user_pk):
    user = request.user
    if user is not None and user.is_authenticated and user.pk == user_pk:

        event_list = []
        if user.is_admin:
            for event in Event.objects.all():
                to_send = {
                    "event" : event,
                }
                event_list.append(to_send)
        else:
            for event in user.event_register_set.all():
                to_send = {
                    "event" : Event.objects.get(pk=event.event_pk),
                    "register" : event,
                }
                event_list.append(to_send)

        response = {
            "user" : user,
            "event_list" : event_list,
        }
        return render(request, "joseph_app/cabinet.html", response)
    else:
        return HttpResponsePermanentRedirect(reverse("joseph_app:index"))

@login_required(login_url="/joseph")
def polls(request):
    user = request.user
    choice_list = []
    for choice in user.user_poll_choice_set.all():
        choice_list.append(choice.poll)
    response = {
        "polls_list" : Poll.objects.all(),
        "choice_list" : choice_list,
        "user" : user,
    }
    return render(request, "joseph_app/polls.html", response)

@login_required(login_url="/joseph")
def poll_create(request, choice_number):
    new_poll = Poll(text=request.POST['text'], pub_date=datetime.datetime.now(), poll_type=request.POST['poll_type'])
    new_poll.save()
    if request.FILES:
        f = request.FILES['poll_image']
        path = STATIC_PATH + FILE_PATH + "poll_img/{}/".format(new_poll.pk)
        if not os.path.exists(path):
            os.mkdir(path[:-1])
        with open(path+f.name, "wb+") as destination:
            for chunk in f.chunks():
                destination.write(chunk)
        new_poll.poll_image = path[len(STATIC_PATH):]+f.name
    new_poll.save()
    for i in range(1,choice_number+1):
        try:
            new_choice = Poll_choice(poll=new_poll, text=request.POST['poll'+str(i)])
            new_choice.save()
        except:
            break
    return HttpResponsePermanentRedirect(reverse("joseph_app:polls"))

@login_required(login_url="/joseph")
def poll_choice_reg(request, poll_pk):
    user = request.user
    poll = Poll.objects.get(pk=poll_pk)
    if poll.poll_type == 1:
        choice = Poll_choice.objects.get(pk=request.POST[str(poll_pk)+'choice'])
        choice.votes += 1
        choice.save()
        user_choice = User_poll_choice(user=user, poll=poll.pk, choice=choice.pk)
        user_choice.save()
    elif poll.poll_type == 2:
        choice_list = request.POST.getlist(str(poll_pk)+'choice')
        for choice_pk in choice_list:
            choice = Poll_choice.objects.get(pk=choice_pk)
            choice.votes += 1
            choice.save()
        user_choice = User_poll_choice(user=user, poll=poll.pk, choice_mult="_".join(choice_list))
        user_choice.save()

    return HttpResponsePermanentRedirect(reverse("joseph_app:polls"))

@login_required(login_url="/joseph")
def events(request):
    user = request.user
    event_list = Event.objects.all().order_by("-event_date")
    badge_list = []
    for event in event_list:
        if event.event_date < timezone.now():
            to_send = {
                "event" : event,
                "badge" : -1,
            }
        if user is not None and not user.is_admin:
            try:
                reg = Event_register.objects.get(user=user, event_pk=event.pk)
                to_send = {
                    "event": event,
                    "badge": reg.has_visited,
                    "reg": reg,
                }
                if event.event_date < timezone.now() and reg.has_visited == False:
                    to_send = {
                        "event": event,
                        "badge": -1,
                    }
            except:
                to_send = {
                    "event": event,
                    "badge": 2,
                }
                if event.event_date < timezone.now():
                    to_send = {
                        "event": event,
                        "badge": -1,
                    }
        else:
            to_send = {
                "event": event,
                "badge": 3,
            }
            if event.event_date < timezone.now():
                to_send = {
                    "event": event,
                    "badge": -1,
                }
        badge_list.append(to_send)
    response = {
        "badge_list" : badge_list,
        "user" : user,
    }
    return render(request, "joseph_app/events.html", response)

@login_required(login_url="/joseph")
def event_create(request):
    user = request.user

    new_event = Event(title=request.POST['event_title'], text=request.POST['event_text'], pub_date=datetime.datetime.now(),
                      event_date=request.POST['event_date'], place=request.POST['place'])
    new_event.save()
    if request.FILES:
        f = request.FILES['event_image']
        path = STATIC_PATH + FILE_PATH + "event_img/{}/".format(new_event.pk)
        if not os.path.exists(path):
            os.mkdir(path[:-1])
        with open(path+f.name, "wb+") as destination:
            for chunk in f.chunks():
                destination.write(chunk)
        new_event.image = path[len(STATIC_PATH):]+f.name
    new_event.save()

    payload = {

         'platform': 'vk',
         'users': 'everyone',
         'data': {
                  'title': Event.objects.get(pk=new_event.pk).title,
                  'text' : Event.objects.get(pk=new_event.pk).text,
                  # 'pub_date' : Event.objects.get(pk=new_event.pk).pub_date,
                  'event_date' : Event.objects.get(pk=new_event.pk).event_date.strftime("%Y-%m-%d %H:%M:%S"),
                  'place' : Event.objects.get(pk=new_event.pk).place,
                  # 'image': open(STATIC_PATH+Event.objects.get(pk=new_event.pk).image, "rb+"),
                 }
              }
    headers = {'content-type': 'application/json'}

    link = 'https://app.botmother.com/api/bot/action/ByqHjGiSX/BFCTDA0D8CsB1WDFWr7CsD5BlBkDXjLBoCImCwrCyBGBp_BLDCDIBODoBHCoClB3'

    r = requests.post(link, data = json.dumps(payload), headers = headers)

    return HttpResponsePermanentRedirect(reverse("joseph_app:user_cab", args=(user.pk,)))

@login_required(login_url="/joseph")
def event_register(request, event_pk):
    user = request.user
    reg_token = Event_register(user=user, event_pk=event_pk)
    reg_token.save()
    link = 'events/visited/' + str(reg_token.pk)
    qr_form = pyqrcode.create(link)
    qr_form.png(STATIC_PATH + FILE_PATH + 'qrcodes/link-{}.png'.format(str(reg_token.pk)), scale=5,  module_color=[0, 0, 0],
        background=[255,255,255])
    reg_token.qr = FILE_PATH + 'qrcodes/link-{}.png'.format(str(reg_token.pk))
    reg_token.save()

    return HttpResponsePermanentRedirect(reverse("joseph_app:events"))

@csrf_exempt
def event_visited(request, reg_pk):
    reg = Event_register.objects.get(pk=reg_pk)
    reg.has_visited = True
    reg.save()

    return HttpResponse(reg.user.name + ' was registrated')

@login_required(login_url="/joseph")
def article_create(request):
    user = request.user
    new_article = Article(title=request.POST['title'], body=request.POST['body'], date=datetime.datetime.now())
    new_article.save()
    if request.FILES:
        for key in request.FILES.keys():
            f = request.FILES[key]
            path = STATIC_PATH+FILE_PATH+"article_gal/{}/".format(new_article.pk)
            if not os.path.exists(path):
                os.mkdir(path[:-1])
            with open(path+f.name, "wb+") as destination:
                for chunk in f.chunks():
                    destination.write(chunk)
            new_gal_img = Article_Image(article=new_article, image=path[len(STATIC_PATH):]+f.name)
            new_gal_img.save()

    return HttpResponsePermanentRedirect(reverse("joseph_app:user_cab"), args=(user.pk,))

def make_docx(request, event_pk):
    document = docxdoc()
    document.add_heading(Event.objects.get(pk=event_pk).title, 0)
    table = document.add_table(rows=1, cols=8)
    hd_row = table.rows[0].cells
    hd_row[0].text = "№"
    hd_row[1].text = "Имя"
    hd_row[2].text = "Фамилия"
    hd_row[3].text = "Электронная почта"
    hd_row[4].text = "Телефон"
    hd_row[5].text = "Вуз"
    hd_row[6].text = "Курс"
    hd_row[7].text = "Посещение"
    i = 0
    for reg in Event_register.objects.filter(event_pk=event_pk):
        data_row = table.add_row().cells
        data_row[0].text = str(i+1)
        data_row[1].text = reg.user.name
        data_row[2].text = reg.user.surname
        data_row[3].text = reg.user.email
        data_row[4].text = reg.user.phone
        data_row[5].text = reg.user.university
        data_row[6].text = str(reg.user.course)
        if reg.has_visited:
            data_row[7].text = "Посетил"
        else:
            data_row[7].text = "Отсутствовал"
        i += 1
    document.add_page_break()
    f = open("temp", mode="wb+")
    document.save(f.name)
    response = HttpResponse(f, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response['content-disposition'] = "attachment; filename='event_brief_{}.docx'".format(event_pk)
    f.close()
    os.remove("./temp")
    return response

def make_xlsx(request, event_pk):
    wb = openpyxl.Workbook()
    wb.active.title = Event.objects.get(pk=event_pk).title
    tab = wb.get_sheet_by_name(wb.active.title)
    h_font = Font(size = 20)
    for l in "ABCDEFGH":
        tab[l+"1"].font = h_font
        tab.column_dimensions[l].width = 30
    tab["A1"] = "№"
    tab["B1"] = "Имя"
    tab["C1"] = "Фамилия"
    tab["D1"] = "Электронная почта"
    tab["E1"] = "Телефон"
    tab["F1"] = "Вуз"
    tab["G1"] = "Курс"
    tab["H1"] = "Посещение"
    i = 2
    for reg in Event_register.objects.filter(event_pk=event_pk):
        tab["A"+str(i)] = str(i-1)
        tab["B"+str(i)] = reg.user.name
        tab["C"+str(i)] = reg.user.surname
        tab["D"+str(i)] = reg.user.email
        tab["E"+str(i)] = reg.user.phone
        tab["F"+str(i)] = reg.user.university
        tab["G"+str(i)] = str(reg.user.course)
        if reg.has_visited:
            tab["H"+str(i)] = "Посетил"
        else:
            tab["H"+str(i)] = "Отсутствовал"
        i += 1
    f = open("temp", mode="wb+")
    wb.save(f.name)
    response = HttpResponse(f, content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    response['content-disposition'] = "attachment; filename='event_brief_{}.xlsx'".format(event_pk)
    f.close()
    os.remove("./temp")
    return response

def retrieve_event_list(request):
    event_list = []
    for event in Event.objects.all():
        to_send = {
            "body" : "{0}&|&{1}&|&{2}&|&{3}&|&{4}".format(str(event.pk), event.title, event.text, event.event_date.strftime("%Y-%m-%d %H:%M:%S"), event.place)
        }
        event_list.append(to_send)
    response = {
        "event_list" : event_list,
    }
    #return render(request, "temp.html", response) #(DBG)
    return JsonResponse(response)

def retrieve_reg_list(request, event_pk):
    reg_obj_list = Event_register.objects.filter(event_pk=event_pk)
    reg_list = []
    for reg in reg_obj_list:
        to_send = {

            "body" : "{0}&|&{1}&|&{2}&|&{3}&|&{4}&|&{5}&|&{6}&|&{7}".format(str(reg.pk), reg.user.email, reg.user.name, reg.user.surname, reg.user.university, str(reg.user.course), reg.user.phone, reg.has_visited)
        }
        reg_list.append(to_send)
    response = {
        "reg_list" : reg_list,
    }
    #return render(request, "temp.html", response) #(DBG)
    return JsonResponse(response)

@login_required(login_url="/joseph")
def file_upload_page(request):
    response = {
        'documents' : Document.objects.all()
    }
    return  render(request, 'joseph_app/docs.html', response)

@login_required(login_url="/joseph")
def file_upload(request):
    doc = Document(title = request.POST['title'], pub_date = datetime.datetime.now())
    doc.save()
    if request.FILES:
        f = request.FILES['document']
        path = STATIC_PATH+FILE_PATH+"documents/"
        with open(path+f.name, "wb+") as destination:
            for chunk in f.chunks():
                destination.write(chunk)

        doc.path = path[len(STATIC_PATH):]+f.name
        doc.size = os.path.getsize(STATIC_PATH + doc.path)
    doc.save()

    return HttpResponsePermanentRedirect(reverse("joseph_app:file_upload_page"))

@login_required(login_url="/joseph")
def file_download(request,doc_pk):
    import mimetypes
    mimetypes.init()
    fname = Document.objects.get(pk=doc_pk).path
    mtg = mimetypes.guess_type(STATIC_PATH + fname)
    f = open(STATIC_PATH + fname, "rb", encoding=mtg[1])
    response = HttpResponse(f, content_type=mtg[0])
    response["Content-Disposition"] = "attachment; filename="+f.name.split('documents')[1]

    return response

# Create your views here.
